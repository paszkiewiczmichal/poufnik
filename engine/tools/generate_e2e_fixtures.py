"""Generate binary fixtures used by desktop WebDriver tests."""

from __future__ import annotations

from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = REPO_ROOT / "corpus" / "e2e"


def main() -> None:
    from docx import Document

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Umowa testowa E2E", level=1)
    document.add_paragraph(
        "Stroną umowy jest Jan Kowalski. PESEL 44051401359. "
        "Sygnatura ABC-XYZ-77. E-mail jan.kowalski@example.com."
    )
    document.save(OUTPUT_DIR / "umowa-e2e.docx")


if __name__ == "__main__":
    main()
