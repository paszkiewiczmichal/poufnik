"""Export anonymized text back to simple document formats."""

from __future__ import annotations

import io
from importlib.resources import as_file, files
from xml.sax.saxutils import escape

from anonymizer_engine.parsers import Block, ParsedDocument


def export_txt(_parsed_doc: ParsedDocument, anonymized_text: str) -> bytes:
    return anonymized_text.encode("utf-8")


def export_docx(parsed_doc: ParsedDocument, anonymized_text: str) -> bytes:
    """Export blocks as a DOCX document.

    Consecutive ``table_cell`` blocks are reconstructed as a table when the original
    tab/newline separators make the row shape unambiguous. Otherwise cells fall back to
    ordinary paragraphs, which preserves text content without pretending to know layout.
    """
    from docx import Document
    from docx.enum.text import WD_BREAK

    document = Document()
    projected = _project_blocks(parsed_doc, anonymized_text)
    index = 0

    while index < len(projected):
        block, content = projected[index]
        if block.kind == "page_break":
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            index += 1
            continue

        if block.kind == "table_cell":
            next_index = _append_table_or_paragraphs(document, parsed_doc, projected, index)
            index = next_index
            continue

        if block.kind == "heading":
            document.add_heading(content, level=1)
        else:
            document.add_paragraph(content)
        index += 1

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_pdf(parsed_doc: ParsedDocument, anonymized_text: str) -> bytes:
    """Export anonymized text to a simple wrapped PDF with embedded DejaVuSans."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    font_name = _register_dejavu_font()
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=48,
        rightMargin=48,
        topMargin=48,
        bottomMargin=48,
    )
    style = ParagraphStyle(
        "AnonymizerBody",
        fontName=font_name,
        fontSize=10,
        leading=14,
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "AnonymizerHeading",
        parent=style,
        fontSize=14,
        leading=18,
        spaceAfter=10,
    )

    story = []
    for block, content in _project_blocks(parsed_doc, anonymized_text):
        if block.kind == "page_break":
            story.append(PageBreak())
            continue
        if not content:
            continue
        paragraph_style = heading_style if block.kind == "heading" else style
        story.append(Paragraph(escape(content).replace("\n", "<br/>"), paragraph_style))
        story.append(Spacer(1, 2))

    if not story:
        story.append(Paragraph("", style))
    document.build(story)
    return buffer.getvalue()


def _append_table_or_paragraphs(
    document: object,
    parsed_doc: ParsedDocument,
    projected: list[tuple[Block, str]],
    start: int,
) -> int:
    run: list[tuple[int, Block, str]] = []
    index = start
    while index < len(projected) and projected[index][0].kind == "table_cell":
        block, content = projected[index]
        run.append((index, block, content))
        index += 1

    rows = _table_rows(parsed_doc, run)
    if rows and len({len(row) for row in rows}) == 1:
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        for row_index, row in enumerate(rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value
        return index

    for _item_index, _block, content in run:
        document.add_paragraph(content)
    return index


def _table_rows(parsed_doc: ParsedDocument, run: list[tuple[int, Block, str]]) -> list[list[str]]:
    rows: list[list[str]] = [[]]
    for run_index, (_projected_index, block, content) in enumerate(run):
        rows[-1].append(content)
        next_is_table_cell = run_index + 1 < len(run)
        if not next_is_table_cell:
            break
        next_block = run[run_index + 1][1]
        gap = parsed_doc.text[block.end : next_block.start]
        if "\n" in gap:
            rows.append([])
    return [row for row in rows if row]


def _project_blocks(parsed_doc: ParsedDocument, anonymized_text: str) -> list[tuple[Block, str]]:
    blocks = parsed_doc.blocks
    if not blocks:
        return []

    projected: list[tuple[Block, str]] = []
    cursor = 0
    leading_gap = parsed_doc.text[: blocks[0].start]
    if leading_gap and anonymized_text.startswith(leading_gap):
        cursor = len(leading_gap)

    for index, block in enumerate(blocks):
        next_start = blocks[index + 1].start if index + 1 < len(blocks) else len(parsed_doc.text)
        gap_after = parsed_doc.text[block.end : next_start]
        if gap_after:
            gap_position = anonymized_text.find(gap_after, cursor)
            if gap_position == -1:
                content = anonymized_text[cursor:] if index == len(blocks) - 1 else ""
                cursor = len(anonymized_text)
            else:
                content = anonymized_text[cursor:gap_position]
                cursor = gap_position + len(gap_after)
        elif index == len(blocks) - 1:
            content = anonymized_text[cursor:]
            cursor = len(anonymized_text)
        else:
            span_length = block.end - block.start
            content = anonymized_text[cursor : cursor + span_length]
            cursor += span_length
        projected.append((block, content))

    return projected


def _register_dejavu_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = "AnonymizerDejaVuSans"
    if font_name in pdfmetrics.getRegisteredFontNames():
        return font_name

    font_resource = files("anonymizer_engine.anonymize.resources").joinpath("DejaVuSans.ttf")
    with as_file(font_resource) as font_path:
        pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
    return font_name
