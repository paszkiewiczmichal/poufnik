from __future__ import annotations

import json
import os
import socket
import subprocess
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import httpx
import pytest
from fastapi.testclient import TestClient

from anonymizer_engine.api import create_app
from anonymizer_engine.api.config import ensure_api_key, rotate_api_key
from anonymizer_engine.api.update_check import UpdateChecker
from anonymizer_engine.detection import (
    DetectedEntity,
    EntityCategory,
    ValidationStatus,
)
from anonymizer_engine.licensing import Plan

API_KEY = "test-api-key"
HEADERS = {"X-Api-Key": API_KEY}


@dataclass(frozen=True)
class Token:
    text: str
    idx: int
    lemma_: str


class NoopNerEngine:
    def __init__(self) -> None:
        self.last_tokens: list[Token] = []

    def analyze(self, text: str, language: str) -> list[DetectedEntity]:
        assert language == "pl"
        return []


@pytest.fixture
def client(tmp_path: Path) -> TestClient:
    app = create_app(api_key=API_KEY, config_dir=tmp_path, ner_engine=NoopNerEngine())
    return TestClient(app)


def test_auth_rejects_missing_and_wrong_api_key(client: TestClient) -> None:
    missing = client.get("/v1/health")
    wrong = client.get("/v1/health", headers={"X-Api-Key": "wrong"})

    assert missing.status_code == 401
    assert wrong.status_code == 401
    assert missing.headers["content-type"].startswith("application/problem+json")


def test_desktop_cors_preflight_allows_tauri_origin(client: TestClient) -> None:
    response = client.options(
        "/v1/documents/process",
        headers={
            "Origin": "http://tauri.localhost",
            "Access-Control-Request-Method": "POST",
            "Access-Control-Request-Headers": "content-type,x-api-key",
        },
    )

    assert response.status_code == 200
    assert response.headers["access-control-allow-origin"] == "http://tauri.localhost"
    assert "x-api-key" in response.headers["access-control-allow-headers"].lower()


def test_serve_auth_reloads_rotated_api_key(tmp_path: Path) -> None:
    old_key, _path, _created = ensure_api_key(tmp_path)
    app = create_app(
        api_key=old_key,
        config_dir=tmp_path,
        reload_api_key_from_config=True,
        ner_engine=NoopNerEngine(),
    )
    client = TestClient(app)

    assert client.get("/v1/health", headers={"X-Api-Key": old_key}).status_code == 200

    new_key, _path, had_previous_key = rotate_api_key(tmp_path)

    assert had_previous_key is True
    assert new_key != old_key
    assert client.get("/v1/health", headers={"X-Api-Key": old_key}).status_code == 401
    assert client.get("/v1/health", headers={"X-Api-Key": new_key}).status_code == 200


def test_health_info_prompts_and_usage_contract(client: TestClient) -> None:
    health = client.get("/v1/health", headers=HEADERS)
    info = client.get("/v1/info", headers=HEADERS)
    prompts = client.get("/v1/prompts", headers=HEADERS)
    usage = client.get("/v1/usage", headers=HEADERS)

    assert health.status_code == 200
    assert health.json()["status"] == "ok"
    assert health.json()["models_loaded"] is True
    assert info.status_code == 200
    assert set(info.json()["categories"]) == {category.value for category in EntityCategory}
    assert info.json()["plan"]["code"] == "basic"
    assert info.json()["update_available"] is False
    assert info.json()["latest_version"] is None
    assert info.json()["degraded_reason"] is None
    assert prompts.status_code == 200
    prompt_id = prompts.json()[0]["id"]
    prompt = client.get(f"/v1/prompts/{prompt_id}", headers=HEADERS)
    assert prompt.status_code == 200
    assert "{{DOKUMENT}}" in prompt.json()["body"]
    assert usage.status_code == 200
    assert usage.json()["requests"] >= 3


def test_analyze_anonymize_and_deanonymize_flow(client: TestClient) -> None:
    text = "Jan ma PESEL 44051401359 i e-mail jan@example.com."

    analyze = client.post("/v1/analyze", headers=HEADERS, json={"text": text})
    assert analyze.status_code == 200
    categories = {entity["category"] for entity in analyze.json()["entities"]}
    assert {"PESEL", "EMAIL"}.issubset(categories)

    anonymized = client.post("/v1/anonymize", headers=HEADERS, json={"text": text})
    assert anonymized.status_code == 200
    payload = anonymized.json()
    assert "[PESEL_1]" in payload["anonymized_text"]
    assert "[EMAIL_1]" in payload["anonymized_text"]
    assert payload["offset_map"]
    assert payload["offset_map"][0]["token"] in payload["anonymized_text"]
    assert {
        "original_start",
        "original_end",
        "anonymized_start",
        "anonymized_end",
        "token",
        "category",
    }.issubset(payload["offset_map"][0])

    deanonymized = client.post(
        "/v1/deanonymize",
        headers=HEADERS,
        json={
            "text": payload["anonymized_text"],
            "replacement_map": payload["replacement_map"],
        },
    )
    assert deanonymized.status_code == 200
    assert deanonymized.json()["original_text"] == text
    assert deanonymized.json()["warnings"] == []


def test_anonymize_accepts_explicit_entities(client: TestClient) -> None:
    text = "Kontakt jan@example.com."
    entity = DetectedEntity(
        category=EntityCategory.EMAIL,
        start=text.index("jan@example.com"),
        end=text.index("jan@example.com") + len("jan@example.com"),
        text="jan@example.com",
        confidence=1.0,
        source="regex",
        validation=ValidationStatus.NOT_APPLICABLE,
    )

    response = client.post(
        "/v1/anonymize",
        headers=HEADERS,
        json={"text": text, "entities": [entity.model_dump(mode="json")]},
    )

    assert response.status_code == 200
    assert response.json()["anonymized_text"] == "Kontakt [EMAIL_1]."


def test_anonymize_accepts_manual_custom_and_rejected_entities(client: TestClient) -> None:
    text = "Numer klienta ABC-123 i jawny fragment."
    accepted = {
        "category": "CUSTOM",
        "start": text.index("ABC-123"),
        "end": text.index("ABC-123") + len("ABC-123"),
        "text": "ABC-123",
        "confidence": 1.0,
        "source": "manual",
        "validation": "not_applicable",
        "status": "accepted",
    }
    rejected = {
        "category": "CUSTOM",
        "start": text.index("jawny"),
        "end": text.index("jawny") + len("jawny"),
        "text": "jawny",
        "confidence": 1.0,
        "source": "manual",
        "validation": "not_applicable",
        "status": "rejected",
    }

    response = client.post(
        "/v1/anonymize",
        headers=HEADERS,
        json={"text": text, "entities": [accepted, rejected]},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["anonymized_text"] == "Numer klienta [DANE_1] i jawny fragment."
    assert payload["replacement_map"]["entries"][0]["category"] == "CUSTOM"


def test_analyze_applies_custom_regex_rules(client: TestClient) -> None:
    text = "Sygnatura klienta ABC-XYZ-77."

    response = client.post(
        "/v1/analyze",
        headers=HEADERS,
        json={
            "text": text,
            "custom_rules": [
                {
                    "name": "Sygnatura",
                    "label": "Sygnatura klienta",
                    "pattern": r"ABC-[A-Z]{3}-\d{2}",
                }
            ],
        },
    )

    assert response.status_code == 200
    payload = response.json()
    custom = [entity for entity in payload["entities"] if entity["category"] == "CUSTOM"]
    assert len(custom) == 1
    assert custom[0]["text"] == "ABC-XYZ-77"
    assert custom[0]["source"] == "regex"


def test_analyze_returns_dictionary_source_for_person(client: TestClient) -> None:
    text = "Dokument podpisał Jan Kowalski."

    response = client.post("/v1/analyze", headers=HEADERS, json={"text": text})

    assert response.status_code == 200
    people = [entity for entity in response.json()["entities"] if entity["category"] == "PERSON"]
    assert len(people) == 1
    assert people[0]["text"] == "Jan Kowalski"
    assert people[0]["source"] == "dictionary"
    assert people[0]["corroborated_by"] == []


def test_analyze_rejects_invalid_custom_regex(client: TestClient) -> None:
    response = client.post(
        "/v1/analyze",
        headers=HEADERS,
        json={
            "text": "ABC",
            "custom_rules": [{"name": "Zła", "label": "Zła", "pattern": "("}],
        },
    )

    assert response.status_code == 422
    assert "pattern is invalid" in str(response.json())


def test_export_docx_and_pdf_contract(client: TestClient) -> None:
    text = "Umowa [OSOBA_1]\n\nPESEL [PESEL_1]"
    blocks = [
        {"start": 0, "end": 15, "kind": "paragraph", "page": 1},
        {"start": 17, "end": len(text), "kind": "paragraph", "page": 1},
    ]

    docx = client.post(
        "/v1/export",
        headers=HEADERS,
        json={"anonymized_text": text, "format": "docx", "blocks": blocks},
    )
    pdf = client.post(
        "/v1/export",
        headers=HEADERS,
        json={"anonymized_text": text, "format": "pdf", "blocks": blocks},
    )

    assert docx.status_code == 200
    assert docx.headers["content-type"].startswith("application/vnd.openxmlformats-officedocument")
    assert docx.content[:2] == b"PK"
    assert pdf.status_code == 200
    assert pdf.headers["content-type"].startswith("application/pdf")
    assert pdf.content.startswith(b"%PDF")


def test_documents_process_contract_for_docx(client: TestClient, tmp_path: Path) -> None:
    docx_path = _write_docx(tmp_path / "umowa.docx")

    response = client.post(
        "/v1/documents/process",
        headers=HEADERS,
        files={
            "file": (
                docx_path.name,
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"force_ocr": "false", "language": "pl"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["document"]["filename"] == "umowa.docx"
    assert payload["document"]["format"] == "docx"
    assert payload["document"]["source"] == "parsed"
    assert "44051401359" in payload["document"]["text"]
    assert "[PESEL_1]" in payload["anonymized_text"]
    assert payload["replacement_map"]["entries"]
    assert payload["offset_map"]


def test_text_size_limit_returns_413_problem_json(client: TestClient) -> None:
    response = client.post(
        "/v1/analyze",
        headers=HEADERS,
        json={"text": "x" * (10 * 1024 * 1024 + 1)},
    )

    assert response.status_code == 413
    assert response.headers["content-type"].startswith("application/problem+json")


def test_plan_limit_enforces_429(tmp_path: Path) -> None:
    plan = Plan(
        code="test",
        monthly_document_limit=None,
        monthly_request_limit=2,
        max_instances=None,
        features=[],
    )
    app = create_app(
        api_key=API_KEY,
        config_dir=tmp_path,
        plan=plan,
        ner_engine=NoopNerEngine(),
    )
    client = TestClient(app)

    assert client.get("/v1/health", headers=HEADERS).status_code == 200
    assert client.get("/v1/health", headers=HEADERS).status_code == 200
    response = client.get("/v1/health", headers=HEADERS)

    assert response.status_code == 429
    assert response.headers["content-type"].startswith("application/problem+json")


def test_start_without_ner_model_fails_without_degraded_flag(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.delenv("ANONYMIZER_ALLOW_DEGRADED", raising=False)
    monkeypatch.setattr("anonymizer_engine.api.app._default_ner_package_available", lambda: False)

    with pytest.raises(RuntimeError, match="Cannot start Anonymizer Engine"):
        create_app(api_key=API_KEY, config_dir=tmp_path)


def test_start_without_ner_model_allows_degraded_with_flag(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr("anonymizer_engine.api.app._default_ner_package_available", lambda: False)
    app = create_app(api_key=API_KEY, config_dir=tmp_path, allow_degraded=True)
    client = TestClient(app)

    health = client.get("/v1/health", headers=HEADERS)
    info = client.get("/v1/info", headers=HEADERS)
    response = client.post(
        "/v1/analyze",
        headers=HEADERS,
        json={"text": "PESEL 44051401359"},
    )

    assert health.status_code == 200
    assert health.json()["status"] == "degraded"
    assert health.json()["models_loaded"] is False
    assert info.status_code == 200
    assert "deterministic and dictionary detection only" in info.json()["degraded_reason"]
    assert response.status_code == 200
    assert response.json()["entities"][0]["category"] == "PESEL"


def test_start_without_ner_model_allows_degraded_with_env(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("ANONYMIZER_ALLOW_DEGRADED", "true")
    monkeypatch.setattr("anonymizer_engine.api.app._default_ner_package_available", lambda: False)

    app = create_app(api_key=API_KEY, config_dir=tmp_path)
    client = TestClient(app)

    response = client.get("/v1/health", headers=HEADERS)

    assert response.status_code == 200
    assert response.json()["status"] == "degraded"
    assert response.json()["models_loaded"] is False


def test_explicit_ner_engine_reports_ok_even_when_default_model_is_unavailable(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr("anonymizer_engine.api.app._default_ner_package_available", lambda: False)

    app = create_app(api_key=API_KEY, config_dir=tmp_path, ner_engine=NoopNerEngine())
    client = TestClient(app)

    health = client.get("/v1/health", headers=HEADERS)
    info = client.get("/v1/info", headers=HEADERS)

    assert health.status_code == 200
    assert health.json()["status"] == "ok"
    assert health.json()["models_loaded"] is True
    assert info.json()["degraded_reason"] is None


def test_documents_process_does_not_open_outbound_connections(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    real_socket = socket.socket

    class GuardedSocket(real_socket):  # type: ignore[misc]
        def connect(self, address: Any) -> None:  # type: ignore[override]
            host = address[0] if isinstance(address, tuple) else ""
            if host not in {"127.0.0.1", "localhost", "::1"}:
                raise AssertionError(f"Outbound connection attempted: {address!r}")
            return super().connect(address)

    monkeypatch.setattr(socket, "socket", GuardedSocket)
    app = create_app(api_key=API_KEY, config_dir=tmp_path, ner_engine=NoopNerEngine())
    client = TestClient(app)
    docx_path = _write_docx(tmp_path / "offline.docx")

    response = client.post(
        "/v1/documents/process",
        headers=HEADERS,
        files={"file": (docx_path.name, docx_path.read_bytes())},
    )

    assert response.status_code == 200


def test_update_check_disabled_by_default_does_not_fetch() -> None:
    calls: list[str] = []

    def fetcher(url: str, _timeout: float) -> str | None:
        calls.append(url)
        raise AssertionError("disabled update check should not fetch")

    checker = UpdateChecker(
        current_version="0.1.0",
        enabled=False,
        manifest_url="https://updates.example/latest",
        fetcher=fetcher,
    )

    checker.start()
    checker.check_once()

    assert calls == []
    assert checker.snapshot().update_available is False
    assert checker.snapshot().latest_version is None


def test_default_update_check_does_not_open_outbound_connections_on_startup(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.delenv("ANONYMIZER_UPDATE_CHECK", raising=False)
    real_socket = socket.socket

    class GuardedSocket(real_socket):  # type: ignore[misc]
        def connect(self, address: Any) -> None:  # type: ignore[override]
            host = address[0] if isinstance(address, tuple) else ""
            if host not in {"127.0.0.1", "localhost", "::1"}:
                raise AssertionError(f"Outbound connection attempted: {address!r}")
            return super().connect(address)

    monkeypatch.setattr(socket, "socket", GuardedSocket)
    app = create_app(api_key=API_KEY, config_dir=tmp_path, ner_engine=NoopNerEngine())

    with TestClient(app) as client:
        response = client.post(
            "/v1/analyze",
            headers=HEADERS,
            json={"text": "PESEL 44051401359"},
        )

    assert response.status_code == 200


def test_update_check_sets_info_fields_with_mocked_manifest(tmp_path: Path) -> None:
    def fetcher(url: str, timeout: float) -> str | None:
        assert url == "https://updates.example/latest"
        assert timeout == 5.0
        return "v0.2.0"

    checker = UpdateChecker(
        current_version="0.1.0",
        enabled=True,
        manifest_url="https://updates.example/latest",
        fetcher=fetcher,
    )
    checker.check_once()
    app = create_app(
        api_key=API_KEY,
        config_dir=tmp_path,
        ner_engine=NoopNerEngine(),
        update_checker=checker,
    )
    client = TestClient(app)

    response = client.get("/v1/info", headers=HEADERS)

    assert response.status_code == 200
    assert response.json()["update_available"] is True
    assert response.json()["latest_version"] == "v0.2.0"


def test_sidecar_mode_prints_port_token_and_requires_token(tmp_path: Path) -> None:
    env = os.environ.copy()
    env["ANONYMIZER_CONFIG_DIR"] = str(tmp_path)
    env["ANONYMIZER_ALLOW_DEGRADED"] = "true"
    env["PYTHONUNBUFFERED"] = "1"
    process = subprocess.Popen(
        [sys.executable, "-m", "anonymizer_engine.api.cli", "--mode", "sidecar", "--port", "0"],
        cwd=Path(__file__).resolve().parents[2],
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    try:
        assert process.stdout is not None
        line = process.stdout.readline().strip()
        if not line:
            stderr = process.stderr.read() if process.stderr is not None else ""
            raise AssertionError(f"sidecar did not print startup JSON; stderr={stderr}")
        startup = json.loads(line)

        health_url = f"http://127.0.0.1:{startup['port']}/v1/health"
        authorized = _get_until_response(
            health_url,
            headers={"X-Api-Key": startup["token"]},
            expected_statuses={200},
        )
        unauthorized = _get_until_response(health_url, expected_statuses={401})

        assert unauthorized.status_code == 401
        assert authorized.status_code == 200
        assert authorized.json()["status"] in {"ok", "degraded"}
    finally:
        process.terminate()
        try:
            process.wait(timeout=10)
        except subprocess.TimeoutExpired:
            process.kill()
            process.wait(timeout=10)


def _write_docx(path: Path) -> Path:
    from docx import Document

    document = Document()
    document.add_paragraph("Umowa testowa")
    document.add_paragraph("PESEL 44051401359, e-mail jan@example.com.")
    document.save(path)
    return path


def _get_until_response(
    url: str,
    *,
    headers: dict[str, str] | None = None,
    expected_statuses: set[int],
    timeout_seconds: float = 15.0,
) -> httpx.Response:
    deadline = time.monotonic() + timeout_seconds
    last_error: Exception | None = None
    while time.monotonic() < deadline:
        try:
            response = httpx.get(url, headers=headers, timeout=1)
        except httpx.HTTPError as exc:
            last_error = exc
        else:
            if response.status_code in expected_statuses:
                return response
            last_error = AssertionError(
                f"Unexpected status {response.status_code}: {response.text}"
            )
        time.sleep(0.2)
    raise AssertionError(f"Timed out waiting for {url}") from last_error
