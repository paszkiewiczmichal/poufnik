from __future__ import annotations

import io
import shutil
import socket
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pytest
from fastapi.testclient import TestClient

from anonymizer_engine.api import create_app
from anonymizer_engine.detection import DetectedEntity

API_KEY = "security-test-key"
HEADERS = {"X-Api-Key": API_KEY}


@dataclass(frozen=True)
class Token:
    text: str
    idx: int
    lemma_: str


class NoopNerEngine:
    def __init__(self) -> None:
        self.last_tokens: list[Token] = []

    def analyze(self, text: str, language: str) -> list[DetectedEntity]:
        assert language == "pl"
        return []


@pytest.fixture
def client(tmp_path: Path) -> TestClient:
    app = create_app(api_key=API_KEY, config_dir=tmp_path, ner_engine=NoopNerEngine())
    return TestClient(app)


def test_analyze_rejects_50_mb_text_with_413(client: TestClient) -> None:
    response = client.post(
        "/v1/analyze",
        headers=HEADERS,
        json={"text": "x" * (50 * 1024 * 1024), "language": "pl"},
    )

    assert response.status_code == 413
    assert response.headers["content-type"].startswith("application/problem+json")
    assert "10 MB" in response.json()["detail"]


def test_analyze_rejects_request_without_content_length(client: TestClient) -> None:
    # A chunked body carries no Content-Length; the size guard must reject it up front so it
    # cannot be buffered unbounded (bypass of the pre-flight size check).
    def body_stream() -> Any:
        yield b'{"text": "a", "language": "pl"}'

    response = client.post(
        "/v1/analyze",
        headers={**HEADERS, "Content-Type": "application/json"},
        content=body_stream(),
    )

    assert response.status_code == 411
    assert response.headers["content-type"].startswith("application/problem+json")


def test_analyze_rejects_oversized_custom_regex_pattern(client: TestClient) -> None:
    response = client.post(
        "/v1/analyze",
        headers=HEADERS,
        json={
            "text": "PESEL 44051401359",
            "language": "pl",
            "custom_rules": [{"name": "evil", "pattern": "(a+)+" * 80}],
        },
    )

    assert response.status_code == 422
    assert response.headers["content-type"].startswith("application/problem+json")
    assert "pattern exceeds" in str(response.json())


def test_analyze_rejects_custom_regex_with_backtracking_risk(client: TestClient) -> None:
    response = client.post(
        "/v1/analyze",
        headers=HEADERS,
        json={
            "text": "a" * 512 + "!",
            "language": "pl",
            "custom_rules": [{"name": "evil", "pattern": r"(a+)+$"}],
        },
    )

    assert response.status_code == 422
    assert response.headers["content-type"].startswith("application/problem+json")
    assert "nested quantifiers" in str(response.json())


def test_detect_custom_rules_times_out_on_heuristic_bypassing_redos() -> None:
    # ``(a|a)*$`` is NOT caught by the nested-quantifier heuristic (no quantifier inside
    # the group), yet with a non-matching tail it is catastrophic (2^n backtracking) under
    # stdlib ``re``. Before the fix this call hung forever; now the hard timeout raises
    # ValueError fast.
    import time

    from anonymizer_engine.detection.deterministic import detect_custom_rules

    started = time.monotonic()
    with pytest.raises(ValueError, match="too slow"):
        detect_custom_rules(
            "a" * 40 + "!",
            [{"name": "evil", "pattern": r"(a|a)*$"}],
            timeout_seconds=0.2,
        )
    assert time.monotonic() - started < 2.0


def test_analyze_redos_poc_pattern_completes_quickly(client: TestClient) -> None:
    # The audit PoC ``(a|a)*b`` hangs the stdlib ``re`` engine; on the ``regex`` engine it
    # resolves in milliseconds. Assert the endpoint returns promptly instead of blocking.
    import time

    started = time.monotonic()
    response = client.post(
        "/v1/analyze",
        headers=HEADERS,
        json={
            "text": "a" * 40 + "X",
            "language": "pl",
            "custom_rules": [{"name": "poc", "pattern": r"(a|a)*b"}],
        },
    )
    elapsed = time.monotonic() - started

    assert response.status_code == 200
    assert elapsed < 2.0


def test_documents_process_rejects_docx_extension_that_is_not_zip(client: TestClient) -> None:
    response = client.post(
        "/v1/documents/process",
        headers=HEADERS,
        files={
            "file": (
                "not-a-zip.docx",
                b"this is not a zip archive",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 400
    assert response.headers["content-type"].startswith("application/problem+json")
    assert "valid ZIP archive" in response.json()["detail"]


def test_documents_process_rejects_docx_zip_bomb_by_uncompressed_size(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import anonymizer_engine.parsers.core as parser_core

    monkeypatch.setattr(parser_core, "MAX_DOCX_UNCOMPRESSED_BYTES", 128)
    payload = io.BytesIO()
    with zipfile.ZipFile(payload, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types></Types>")
        archive.writestr("word/document.xml", "A" * 1024)

    response = client.post(
        "/v1/documents/process",
        headers=HEADERS,
        files={
            "file": (
                "bomb.docx",
                payload.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 400
    assert response.headers["content-type"].startswith("application/problem+json")
    assert "uncompressed size exceeds" in response.json()["detail"]


def test_documents_process_rejects_pdf_with_5000_pages(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import pdfplumber

    class FakePdf:
        pages = [object()] * 5000

        def __enter__(self) -> FakePdf:
            return self

        def __exit__(self, *_args: object) -> None:
            return None

    monkeypatch.setattr(pdfplumber, "open", lambda _source: FakePdf())

    response = client.post(
        "/v1/documents/process",
        headers=HEADERS,
        files={"file": ("too-many-pages.pdf", b"%PDF-1.7\n%", "application/pdf")},
    )

    assert response.status_code == 400
    assert response.headers["content-type"].startswith("application/problem+json")
    assert "PDF page count exceeds" in response.json()["detail"]


def test_documents_process_rejects_force_ocr_pdf_with_5000_pages(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import pypdfium2

    class FakePdf:
        def __init__(self, _data: bytes) -> None:
            self.closed = False

        def __len__(self) -> int:
            return 5000

        def close(self) -> None:
            self.closed = True

    monkeypatch.setattr(pypdfium2, "PdfDocument", FakePdf)

    response = client.post(
        "/v1/documents/process",
        headers=HEADERS,
        files={"file": ("too-many-pages.pdf", b"%PDF-1.7\n%", "application/pdf")},
        data={"force_ocr": "true"},
    )

    assert response.status_code == 400
    assert response.headers["content-type"].startswith("application/problem+json")
    assert "PDF page count exceeds" in response.json()["detail"]


def test_pdf_render_scale_is_clamped_for_extreme_mediabox() -> None:
    from anonymizer_engine.ocr.core import MAX_OCR_PAGE_PIXELS, _render_scale_for_page

    base_scale = 300 / 72
    # A ~200x200 inch page (14400x14400 pt) at 300 DPI would rasterize to ~60000x60000 px
    # (≈3.6 Gpx, multi-GB bitmap). The clamp must bound the output to the safety limit.
    huge = 14400.0
    clamped = _render_scale_for_page(huge, huge, base_scale)
    assert (huge * clamped) * (huge * clamped) <= MAX_OCR_PAGE_PIXELS + 1
    # A normal A4 page (595x842 pt) is left at full resolution.
    assert _render_scale_for_page(595.0, 842.0, base_scale) == base_scale


def test_ocr_image_guard_rejects_oversized_dimensions() -> None:
    from types import SimpleNamespace

    from anonymizer_engine.ocr.core import _ensure_image_within_limits
    from anonymizer_engine.ocr.exceptions import OcrExecutionError

    with pytest.raises(OcrExecutionError, match="exceed the OCR safety limit"):
        _ensure_image_within_limits(SimpleNamespace(size=(50000, 50000)))  # type: ignore[arg-type]
    _ensure_image_within_limits(SimpleNamespace(size=(1000, 1000)))  # type: ignore[arg-type]


def test_ocr_pdf_clamps_render_scale_for_extreme_page(monkeypatch: pytest.MonkeyPatch) -> None:
    import pypdfium2

    from anonymizer_engine.ocr import core as ocr_core

    recorded: dict[str, float] = {}

    class FakeBitmap:
        def to_pil(self) -> Any:
            from PIL import Image as PilImage

            return PilImage.new("RGB", (2, 2))

        def close(self) -> None:
            return None

    class FakePage:
        def get_size(self) -> tuple[float, float]:
            return (14400.0, 14400.0)

        def render(self, scale: float) -> FakeBitmap:
            recorded["scale"] = scale
            return FakeBitmap()

        def close(self) -> None:
            return None

    class FakePdf:
        def __init__(self, _data: bytes) -> None:
            pass

        def __len__(self) -> int:
            return 1

        def __getitem__(self, _index: int) -> FakePage:
            return FakePage()

        def close(self) -> None:
            return None

    monkeypatch.setattr(pypdfium2, "PdfDocument", FakePdf)
    monkeypatch.setattr(ocr_core, "ocr_image", lambda image, languages="pol+eng": "")

    ocr_core.ocr_pdf(b"%PDF-1.7")

    scale = recorded["scale"]
    assert (14400.0 * scale) * (14400.0 * scale) <= ocr_core.MAX_OCR_PAGE_PIXELS + 1


def _real_ner_available() -> bool:
    from anonymizer_engine.detection.ner import spacy_model_available

    return spacy_model_available()


@pytest.mark.skipif(
    not _real_ner_available(),
    reason="pl_core_news_lg model is required to exercise the real NER engine offline.",
)
def test_real_ner_engine_does_not_open_outbound_connections(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # Unlike the Noop-based guard test, this loads and runs the actual SpacyPresidioEngine
    # under the socket guard, proving the shipped detection layer performs no network I/O.
    _install_outbound_socket_guard(monkeypatch)
    app = create_app(api_key=API_KEY, config_dir=tmp_path)
    client = TestClient(app)

    response = client.post(
        "/v1/analyze",
        headers=HEADERS,
        json={"text": "Jan Kowalski mieszka w Warszawie. PESEL 44051401359.", "language": "pl"},
    )

    assert response.status_code == 200
    categories = {item["category"] for item in response.json()["entities"]}
    assert "PESEL" in categories


def test_full_engine_pipeline_does_not_open_outbound_connections(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _install_outbound_socket_guard(monkeypatch)
    app = create_app(api_key=API_KEY, config_dir=tmp_path, ner_engine=NoopNerEngine())
    client = TestClient(app)
    docx_path = _write_docx(tmp_path / "offline.docx")

    processed = client.post(
        "/v1/documents/process",
        headers=HEADERS,
        files={
            "file": (
                docx_path.name,
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert processed.status_code == 200
    assert "[PESEL_1]" in processed.json()["anonymized_text"]

    analyze = client.post(
        "/v1/analyze",
        headers=HEADERS,
        json={"text": "PESEL 44051401359, e-mail jan@example.com."},
    )
    assert analyze.status_code == 200
    assert {item["category"] for item in analyze.json()["entities"]} >= {"PESEL", "EMAIL"}

    anonymized = client.post(
        "/v1/anonymize",
        headers=HEADERS,
        json={"text": "PESEL 44051401359, e-mail jan@example.com."},
    )
    assert anonymized.status_code == 200
    anonymized_payload = anonymized.json()

    docx = client.post(
        "/v1/export",
        headers=HEADERS,
        json={"anonymized_text": anonymized_payload["anonymized_text"], "format": "docx"},
    )
    pdf = client.post(
        "/v1/export",
        headers=HEADERS,
        json={"anonymized_text": anonymized_payload["anonymized_text"], "format": "pdf"},
    )
    assert docx.status_code == 200
    assert docx.content.startswith(b"PK")
    assert pdf.status_code == 200
    assert pdf.content.startswith(b"%PDF")

    restored = client.post(
        "/v1/deanonymize",
        headers=HEADERS,
        json={
            "text": anonymized_payload["anonymized_text"],
            "replacement_map": anonymized_payload["replacement_map"],
        },
    )
    assert restored.status_code == 200
    assert "44051401359" in restored.json()["original_text"]


def _tesseract_supports(languages: str) -> bool:
    command = shutil.which("tesseract")
    if command is None:
        return False
    try:
        result = subprocess.run(
            [command, "--list-langs"],
            check=False,
            capture_output=True,
            text=True,
            timeout=15,
        )
    except (OSError, subprocess.TimeoutExpired):
        return False
    if result.returncode != 0:
        return False
    available = {
        line.strip()
        for line in (result.stdout + "\n" + result.stderr).splitlines()
        if line.strip() and not line.startswith("List of available languages")
    }
    return set(languages.split("+")).issubset(available)


@pytest.mark.skipif(
    not _tesseract_supports("pol+eng"),
    reason="Tesseract with pol+eng language data is required for OCR offline guard.",
)
def test_ocr_document_process_does_not_open_outbound_connections(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _install_outbound_socket_guard(monkeypatch)
    app = create_app(api_key=API_KEY, config_dir=tmp_path, ner_engine=NoopNerEngine())
    client = TestClient(app)

    response = client.post(
        "/v1/documents/process",
        headers=HEADERS,
        files={"file": ("scan.png", _render_ocr_png(), "image/png")},
    )

    assert response.status_code == 200
    assert "44051401359" in response.json()["document"]["text"]
    assert "[PESEL_1]" in response.json()["anonymized_text"]


def _install_outbound_socket_guard(monkeypatch: pytest.MonkeyPatch) -> None:
    real_socket = socket.socket

    class GuardedSocket(real_socket):  # type: ignore[misc]
        def connect(self, address: Any) -> None:  # type: ignore[override]
            host = address[0] if isinstance(address, tuple) else ""
            if host not in {"127.0.0.1", "localhost", "::1"}:
                raise AssertionError(f"Outbound connection attempted: {address!r}")
            return super().connect(address)

    monkeypatch.setattr(socket, "socket", GuardedSocket)


def _write_docx(path: Path) -> Path:
    from docx import Document

    document = Document()
    document.add_paragraph("Umowa testowa")
    document.add_paragraph("Jan Kowalski, PESEL 44051401359, e-mail jan@example.com.")
    document.save(path)
    return path


def _render_ocr_png() -> bytes:
    from PIL import Image, ImageDraw, ImageFont

    text = "PESEL 44051401359"
    font = ImageFont.truetype(str(_find_font_path()), size=64)
    image = Image.new("RGB", (900, 180), "white")
    draw = ImageDraw.Draw(image)
    draw.text((40, 50), text, fill="black", font=font)
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", dpi=(300, 300))
    return buffer.getvalue()


def _find_font_path() -> Path:
    candidates = [
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    pytest.skip("No local TrueType font found for OCR fixture.")
