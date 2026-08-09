from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import pytest

from anonymizer_engine.detection import EntityCategory, detect_all
from anonymizer_engine.detection.models import DetectedEntity
from anonymizer_engine.parsers import (
    CorruptedFile,
    ParsedDocument,
    UnsupportedFormat,
    parse_document,
    parse_pdf,
)


@dataclass(frozen=True)
class Token:
    text: str
    idx: int
    lemma_: str


class NoopNerEngine:
    def __init__(self) -> None:
        self.last_tokens: list[Token] = []

    def analyze(self, text: str, language: str) -> list[DetectedEntity]:
        assert language == "pl"
        return []


def test_parse_txt_detects_cp1250_and_normalizes_newlines(tmp_path: Path) -> None:
    path = tmp_path / "notatka.txt"
    text = "Zażółć gęślą jaźń\r\nciąg dalszy\r\n\r\nPESEL 44051401359"
    path.write_bytes(text.encode("cp1250"))

    parsed = parse_document(path)

    assert parsed.format == "txt"
    assert parsed.has_text_layer is True
    assert parsed.page_count == 1
    assert parsed.text == "Zażółć gęślą jaźń\nciąg dalszy\n\nPESEL 44051401359"
    assert _block_texts(parsed) == [
        "Zażółć gęślą jaźń\nciąg dalszy",
        "PESEL 44051401359",
    ]
    _assert_block_offsets(parsed)


def test_parse_docx_preserves_document_order_tables_and_header_footer(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "fixture_umowa.docx"
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "Nagłówek: Zażółć gęślą jaźń"
    section.footer.paragraphs[0].text = "Stopka: Łódź"
    document.add_heading("Umowa testowa", level=1)
    document.add_paragraph(
        "Umowę zawarto z Janem Kowalskim, PESEL 44051401359, e-mail jan@example.com."
    )
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Pole"
    table.cell(0, 1).text = "Wartość"
    table.cell(1, 0).text = "NIP"
    table.cell(1, 1).text = "526-10-40-828"
    document.save(path)

    parsed = parse_document(path)

    assert parsed.format == "docx"
    assert parsed.has_text_layer is True
    assert parsed.page_count == 1
    assert _block_texts(parsed) == [
        "Nagłówek: Zażółć gęślą jaźń",
        "Umowa testowa",
        "Umowę zawarto z Janem Kowalskim, PESEL 44051401359, e-mail jan@example.com.",
        "Pole",
        "Wartość",
        "NIP",
        "526-10-40-828",
        "Stopka: Łódź",
    ]
    assert [block.kind for block in parsed.blocks] == [
        "paragraph",
        "heading",
        "paragraph",
        "table_cell",
        "table_cell",
        "table_cell",
        "table_cell",
        "paragraph",
    ]
    _assert_block_offsets(parsed)


def test_parse_pdf_extracts_pages_and_page_break_blocks(tmp_path: Path) -> None:
    path = tmp_path / "tekstowy.pdf"
    _write_text_pdf(
        path,
        [
            "Strona 1: Zażółć gęślą jaźń. PESEL 44051401359.",
            "Strona 2: Łódź i Kraków. NIP 526-10-40-828.",
            "Strona 3: Kontakt anna.nowak@example.com.",
        ],
    )

    parsed = parse_pdf(path)

    assert parsed.format == "pdf"
    assert parsed.has_text_layer is True
    assert parsed.page_count == 3
    assert parsed.text.count("\f") == 2
    assert [block.kind for block in parsed.blocks].count("page_break") == 2
    assert "PESEL 44051401359" in parsed.text
    assert "NIP 526-10-40-828" in parsed.text
    assert "anna.nowak@example.com" in parsed.text
    _assert_block_offsets(parsed)


def test_parse_pdf_flags_image_only_document_without_text_layer(tmp_path: Path) -> None:
    path = tmp_path / "skan.pdf"
    _write_image_only_pdf(path, tmp_path / "page.png")

    parsed = parse_pdf(path)

    assert parsed.format == "pdf"
    assert parsed.has_text_layer is False
    assert parsed.page_count == 1
    assert parsed.text == ""
    assert parsed.blocks == []


def test_parse_document_rejects_unknown_extension(tmp_path: Path) -> None:
    path = tmp_path / "plik.xyz"
    path.write_text("Zażółć gęślą jaźń", encoding="utf-8")

    with pytest.raises(UnsupportedFormat):
        parse_document(path)


def test_parse_document_wraps_corrupted_pdf(tmp_path: Path) -> None:
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"%PDF-1.4\nbroken")

    with pytest.raises(CorruptedFile):
        parse_document(path)


def test_parser_output_integrates_with_detection(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "fixture_umowa.docx"
    document = Document()
    document.add_paragraph("Umowę zawarto z Janem Kowalskim.")
    document.add_paragraph("PESEL 44051401359, NIP 526-10-40-828, e-mail jan@example.com.")
    document.save(path)

    result = detect_all(parse_document(path).text, ner_engine=NoopNerEngine())

    categories = {entity.category for entity in result.entities}
    assert EntityCategory.PESEL in categories
    assert EntityCategory.NIP in categories
    assert EntityCategory.EMAIL in categories


def _block_texts(parsed: ParsedDocument) -> list[str]:
    return [parsed.text[block.start : block.end] for block in parsed.blocks]


def _assert_block_offsets(parsed: ParsedDocument) -> None:
    for block in parsed.blocks:
        assert 0 <= block.start <= block.end <= len(parsed.text)
        fragment = parsed.text[block.start : block.end]
        if block.kind == "page_break":
            assert fragment == "\f"
            assert block.page is None
            continue
        assert fragment
        assert not fragment.startswith(("\n", "\t", "\f"))
        assert not fragment.endswith(("\n", "\t", "\f"))


def _write_text_pdf(path: Path, pages: list[str]) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    font_name = _register_pdf_font()
    _width, height = A4
    pdf = canvas.Canvas(str(path), pagesize=A4)
    for page_text in pages:
        pdf.setFont(font_name, 12)
        pdf.drawString(72, height - 72, page_text)
        pdf.drawString(72, height - 92, "Drugi akapit kontrolny z polskimi znakami: ąęłńóśźż.")
        pdf.showPage()
    pdf.save()


def _write_image_only_pdf(path: Path, image_path: Path) -> None:
    from PIL import Image
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    Image.new("RGB", (160, 80), color="white").save(image_path)
    _width, height = A4
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.drawImage(str(image_path), 72, height - 180, width=160, height=80)
    pdf.showPage()
    pdf.save()


def _register_pdf_font() -> str:
    import reportlab
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = "AnonymizerTestFont"
    if font_name in pdfmetrics.getRegisteredFontNames():
        return font_name

    reportlab_dir = Path(reportlab.__file__).resolve().parent
    candidates = [
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        reportlab_dir / "fonts" / "Vera.ttf",
    ]
    for candidate in candidates:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont(font_name, str(candidate)))
            return font_name
    return "Helvetica"
